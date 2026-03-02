from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.cv import CVData, ContactInfo, ExperienceEntry, EducationEntry, SkillGroup
from templates import SECTION_TITLES


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _set_run_font(run, size_pt: float, bold: bool = False, color: tuple = None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_horizontal_rule(doc: Document):
    """Insert a thin bottom-border paragraph as a visual divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    _set_run_font(run, 9, bold=True, color=(90, 90, 90))


# ─── Style Setup ─────────────────────────────────────────────────────────────

def apply_styles(doc: Document, template_id: str = "classic") -> None:
    """Configure global document margins and default paragraph font."""
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)


# ─── Section Renderers ───────────────────────────────────────────────────────

def render_contact(doc: Document, contact: ContactInfo) -> None:
    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(contact.name)
    _set_run_font(name_run, 22, bold=True)

    # Title
    if contact.title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(4)
        title_run = title_p.add_run(contact.title)
        _set_run_font(title_run, 11, color=(80, 80, 80))

    # Contact line
    parts = []
    if contact.location:
        parts.append(contact.location)
    if contact.phone:
        parts.append(contact.phone)
    if contact.email:
        parts.append(contact.email)
    if contact.linkedin:
        parts.append(contact.linkedin)
    if contact.portfolio:
        parts.append(contact.portfolio)

    if parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(6)
        contact_run = contact_p.add_run("  ·  ".join(parts))
        _set_run_font(contact_run, 9, color=(100, 100, 100))


def render_summary(doc: Document, summary: str, title: str) -> None:
    _section_heading(doc, title)
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(10)


def render_skills(doc: Document, skills: List[SkillGroup], title: str) -> None:
    _section_heading(doc, title)
    for group in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        cat_run = p.add_run(f"{group.category}: ")
        _set_run_font(cat_run, 10, bold=True)
        items_run = p.add_run(", ".join(group.items))
        _set_run_font(items_run, 10)


def render_experience(doc: Document, experience: List[ExperienceEntry], title: str) -> None:
    _section_heading(doc, title)
    for entry in experience:
        # Title line: bold job_title + normal " — Company, Location"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        title_run = p.add_run(entry.job_title)
        _set_run_font(title_run, 11, bold=True)
        company_str = f"  —  {entry.company}"
        if entry.location:
            company_str += f", {entry.location}"
        company_run = p.add_run(company_str)
        _set_run_font(company_run, 10, color=(80, 80, 80))

        # Dates
        date_p = doc.add_paragraph(f"{entry.start_date} – {entry.end_date}")
        date_p.paragraph_format.space_after = Pt(2)
        _set_run_font(date_p.runs[0], 9, color=(120, 120, 120))

        # Bullets
        for bullet in entry.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.space_after = Pt(1)
            run = bp.add_run(bullet.text)
            run.font.size = Pt(10)


def render_education(doc: Document, education: List[EducationEntry], title: str) -> None:
    _section_heading(doc, title)
    for entry in education:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        degree_run = p.add_run(entry.degree)
        _set_run_font(degree_run, 10, bold=True)
        inst_run = p.add_run(f"  —  {entry.institution}")
        _set_run_font(inst_run, 10, color=(80, 80, 80))

        date_p = doc.add_paragraph(f"{entry.start_date} – {entry.end_date}")
        date_p.paragraph_format.space_after = Pt(2)
        _set_run_font(date_p.runs[0], 9, color=(120, 120, 120))


# ─── Public API ──────────────────────────────────────────────────────────────

def export_docx(cv: CVData, template_id: str = "classic", language: str = "en") -> bytes:
    """Convert a CVData object to a .docx file and return the raw bytes."""
    titles = SECTION_TITLES.get(language, SECTION_TITLES["en"])
    doc = Document()
    apply_styles(doc, template_id)
    render_contact(doc, cv.contact)
    render_summary(doc, cv.summary, titles["summary"])
    render_skills(doc, cv.skills, titles["skills"])
    render_experience(doc, cv.experience, titles["experience"])
    render_education(doc, cv.education, titles["education"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
